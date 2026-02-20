import wikipedia
from duckduckgo_search import DDGS
from fpdf import FPDF
import datetime
import requests
from bs4 import BeautifulSoup
import os
from PIL import Image
import customtkinter as ctk
import threading
import tkinter.messagebox as msgbox
from tkinter import filedialog
import matplotlib.pyplot as plt
from collections import Counter
import re
import json
from docx import Document
from docx.shared import Inches

# --- CONFIG & UTILS ---
HISTORY_FILE = "search_history.json"

def clean_text(text):
    if not text: return ""
    replacements = {
        '€': 'EUR', '”': '"', '“': '"', '’': "'", '‘': "'", '–': '-', '…': '...',
        'à': 'a', 'è': 'e', 'é': 'e', 'ì': 'i', 'ò': 'o', 'ù': 'u',
        'Á': 'A', 'É': 'E', 'Í': 'I', 'Ó': 'O', 'Ú': 'U',
        'ä': 'a', 'ö': 'o', 'ü': 'u', 'ß': 'ss', 'ñ': 'n'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode('latin-1', 'replace').decode('latin-1')

def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r") as f:
                return json.load(f)
        except:
            return []
    return []

def save_to_history(topic, filepath):
    history = load_history()
    entry = {
        "topic": topic,
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "path": filepath
    }
    # Evita duplicati identici recenti
    if history and history[0]['topic'] == topic and history[0]['path'] == filepath:
        return
    history.insert(0, entry) # Aggiungi in cima
    if len(history) > 20: history = history[:20] # Mantieni solo gli ultimi 20
    
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=4)

def download_image(topic):
    try:
        page = wikipedia.page(topic)
        if page.images:
            for img_url in page.images[:5]:
                if img_url.lower().endswith(('.jpg', '.png', '.jpeg')):
                    response = requests.get(img_url, stream=True, timeout=5)
                    if response.status_code == 200:
                        filename = f"temp_{topic.replace(' ', '')}.jpg"
                        with open(filename, 'wb') as f:
                            f.write(response.content)
                        return filename
    except:
        pass
    return None

def smart_scrape(url):
    try:
        headers = {'User-Agent': "Mozilla/5.0"}
        response = requests.get(url, headers=headers, timeout=4)
        soup = BeautifulSoup(response.content, 'html.parser')
        paragraphs = soup.find_all('p')
        text = " ".join([p.get_text() for p in paragraphs[:3]])
        if len(text) > 400: text = text[:400] + "..."
        return text.strip() or "Nessun contenuto testuale rilevato."
    except:
        return "Accesso al sito non riuscito."

def create_chart(text_data, topic):
    try:
        # Trova parole > 4 lettere
        words = re.findall(r'\b\w{5,}\b', text_data.lower())
        ignore = {'anche', 'della', 'delle', 'nella', 'hanno', 'stato', 'sono', 'come', 'questo', 'questa', 'degli', 'parte', 'prima', 'dopo', 'tutto', 'tutti', 'fatto', 'essere', 'avere', 'which', 'their', 'about', 'would', 'these', 'other', 'sur', 'pour', 'dans', 'avec', 'plus', 'not', 'that', 'with', 'from', 'this', 'have'}
        filtered = [w for w in words if w not in ignore]
        
        counts = Counter(filtered).most_common(7)
        if not counts: return None
        
        labels, values = zip(*counts)
        
        plt.figure(figsize=(6, 4))
        plt.bar(labels, values, color='#213363')
        plt.title(f"Parole Chiave: {topic}")
        plt.xticks(rotation=45)
        plt.tight_layout()
        
        filename = "temp_chart.png"
        plt.savefig(filename)
        plt.close()
        return filename
    except Exception:
        return None

# --- EXPORT LOGIC ---
def generate_docx(topic, wiki_summary, web_results, img_file, chart_file, save_path):
    doc = Document()
    doc.add_heading(f'Report: {topic}', 0)
    
    doc.add_paragraph(f"Generato il: {datetime.date.today()}")

    if img_file:
        try:
            doc.add_picture(img_file, width=Inches(4))
        except: pass

    doc.add_heading('Panoramica Generale', level=1)
    doc.add_paragraph(wiki_summary)

    if chart_file:
        doc.add_heading('Analisi Dati', level=1)
        try:
            doc.add_picture(chart_file, width=Inches(5))
        except: pass

    if web_results:
        doc.add_heading('Risorse Web', level=1)
        for res in web_results:
            p = doc.add_paragraph()
            runner = p.add_run(res['title'])
            runner.bold = True
            doc.add_paragraph(res['body'])
            doc.add_paragraph(res['href'], style='Intense Quote')

    doc.save(save_path)

class PDFReport(FPDF):
    def __init__(self, topic, cover_image=None):
        super().__init__()
        self.topic = clean_text(topic)
        self.cover_image = cover_image
        self.set_auto_page_break(auto=True, margin=15)

    def header(self):
        if self.page_no() > 1:
            self.set_font('Arial', 'I', 8)
            self.set_text_color(128)
            self.cell(0, 10, f'Report: {self.topic}', 0, 0, 'R')
            self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128)
        self.cell(0, 10, f'Pagina {self.page_no()}', 0, 0, 'C')

    def create_cover_page(self):
        self.add_page()
        self.set_y(40)
        self.set_font('Arial', 'B', 24)
        self.set_text_color(33, 51, 99)
        self.cell(0, 20, "REPORT PREMIUM", ln=True, align='C')
        
        if self.cover_image:
            try:
                self.image(self.cover_image, x=65, y=70, w=80)
                self.set_y(160)
            except:
                self.set_y(100)
        else:
            self.set_y(100)

        self.set_font('Arial', 'B', 32)
        self.set_text_color(0)
        self.cell(0, 20, self.topic, ln=True, align='C')
        self.add_page()

    def add_section_title(self, title):
        self.set_font('Arial', 'B', 14)
        self.set_text_color(33, 51, 99)
        self.cell(0, 10, clean_text(title), 0, 1, 'L')
        self.set_text_color(0)
        self.ln(2)

    def add_paragraph(self, text):
        self.set_font('Arial', '', 11)
        self.multi_cell(0, 6, clean_text(text))
        self.ln(5)

    def add_web_card(self, title, body, link):
        self.set_fill_color(240, 240, 240)
        self.set_font('Arial', 'B', 10)
        self.cell(0, 8, clean_text(title), 0, 1, 'L', True)
        self.set_font('Arial', '', 10)
        self.multi_cell(0, 5, clean_text(body))
        self.set_font('Arial', 'I', 8)
        self.set_text_color(0, 0, 255)
        self.cell(0, 5, clean_text(link), 0, 1)
        self.set_text_color(0)
        self.ln(3)

def generate_report(topic, lang, depth, save_path, export_format):
    lang_map = {"Italiano": "it", "English": "en", "Français": "fr", "Español": "es", "Deutsch": "de"}
    code = lang_map.get(lang, "it")
    sentences = 5 if depth == "Veloce" else (20 if depth == "Approfondita" else 10)
    web_count = 1 if depth == "Veloce" else (5 if depth == "Approfondita" else 3)
    
    print(f"Working on {topic} ({export_format})...")
    
    try:
        wikipedia.set_lang(code)
        wiki_summary = wikipedia.summary(topic, sentences=sentences)
        img_file = download_image(topic)
    except:
        wiki_summary = "N/A"
        img_file = None

    web_results = []
    full_text = wiki_summary
    try:
        with DDGS() as ddgs:
            raw = list(ddgs.text(topic, max_results=web_count))
            for r in raw:
                body = smart_scrape(r['href']) if depth != "Veloce" else r['body']
                web_results.append({'title': r['title'], 'body': body, 'href': r['href']})
                full_text += " " + body
    except: pass

    chart_file = create_chart(full_text, topic) if depth != "Veloce" else None

    # Export
    if "PDF" in export_format:
        pdf = PDFReport(topic, img_file)
        pdf.create_cover_page()
        pdf.add_section_title(f"Panoramica ({lang})")
        pdf.add_paragraph(wiki_summary)
        pdf.ln()
        if chart_file:
            pdf.add_section_title("Analisi Semantica")
            pdf.image(chart_file, x=50, w=110)
            pdf.ln(10)
        if web_results:
            pdf.add_section_title("Risorse Web")
            for res in web_results:
                pdf.add_web_card(res['title'], res['body'], res['href'])
        pdf.output(save_path)
    
    elif "Word" in export_format:
        generate_docx(topic, wiki_summary, web_results, img_file, chart_file, save_path)

    # Cleanup & History
    if img_file and os.path.exists(img_file): os.remove(img_file)
    if chart_file and os.path.exists(chart_file): os.remove(chart_file)
    
    save_to_history(topic, save_path)
    return save_path

# --- GUI ---
class UltimateApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Research Station Premium")
        self.geometry("800x650")
        ctk.set_appearance_mode("Dark")
        ctk.set_default_color_theme("blue") 

        # Tabs
        self.tabview = ctk.CTkTabview(self, width=780, height=600)
        self.tabview.pack(pady=10)
        self.tabview.add("Ricerca")
        self.tabview.add("Cronologia")
        self.tabview.add("Impostazioni")
        
        self.setup_search_tab()
        self.setup_history_tab()
        self.setup_settings_tab()

    def setup_search_tab(self):
        tab = self.tabview.tab("Ricerca")
        
        ctk.CTkLabel(tab, text="Research Station", font=("Arial", 28, "bold"), text_color="#4B8BBE").pack(pady=15)
        
        self.entry = ctk.CTkEntry(tab, placeholder_text="Argomento...", width=500, height=45, font=("Arial", 14))
        self.entry.pack(pady=10)

        frame = ctk.CTkFrame(tab, fg_color="transparent")
        frame.pack(pady=5)
        
        ctk.CTkLabel(frame, text="Lingua:").grid(row=0, column=0, padx=5)
        self.opt_lang = ctk.CTkOptionMenu(frame, values=["Italiano", "English", "Français", "Deutsch"])
        self.opt_lang.grid(row=0, column=1, padx=10)
        
        ctk.CTkLabel(frame, text="Profondità:").grid(row=0, column=2, padx=5)
        self.seg_depth = ctk.CTkSegmentedButton(frame, values=["Veloce", "Normale", "Approfondita"])
        self.seg_depth.set("Normale")
        self.seg_depth.grid(row=0, column=3, padx=10)

        ctk.CTkLabel(frame, text="Formato:").grid(row=1, column=0, padx=5, pady=10)
        self.opt_format = ctk.CTkOptionMenu(frame, values=["PDF", "Word (.docx)"])
        self.opt_format.grid(row=1, column=1, padx=10, pady=10)

        self.btn_go = ctk.CTkButton(tab, text="AVVIA RICERCA", width=250, height=50, font=("Arial", 16, "bold"), command=self.ask_save)
        self.btn_go.pack(pady=20)

        self.progress = ctk.CTkProgressBar(tab, width=500, mode="indeterminate")
        self.progress.pack_forget()
        self.status = ctk.CTkLabel(tab, text="Pronto.", text_color="gray")
        self.status.pack(pady=5)
        
        self.btn_open = ctk.CTkButton(tab, text="APRI FILE", fg_color="#E09F3E", state="disabled", command=self.open_current)
        self.btn_open.pack_forget()
        self.saved_path = None

    def setup_history_tab(self):
        tab = self.tabview.tab("Cronologia")
        self.history_frame = ctk.CTkScrollableFrame(tab, width=700, height=450)
        self.history_frame.pack(pady=10)
        self.refresh_history()
        
        ctk.CTkButton(tab, text="Aggiorna Lista", command=self.refresh_history).pack(pady=5)

    def refresh_history(self):
        for w in self.history_frame.winfo_children(): w.destroy()
        hist = load_history()
        if not hist:
            ctk.CTkLabel(self.history_frame, text="Nessuna ricerca recente.").pack(pady=20)
            return
            
        for h in hist:
            f = ctk.CTkFrame(self.history_frame)
            f.pack(fill="x", pady=2, padx=5)
            ctk.CTkLabel(f, text=f"[{h['date']}] {h['topic']}", anchor="w", width=300).pack(side="left", padx=10)
            ctk.CTkButton(f, text="Apri", width=80, command=lambda p=h['path']: self.safe_open(p)).pack(side="right", padx=10)

    def safe_open(self, path):
        if os.path.exists(path):
            os.startfile(path)
        else:
            msgbox.showerror("Errore", "Il file non esiste più.")

    def setup_settings_tab(self):
        tab = self.tabview.tab("Impostazioni")
        ctk.CTkLabel(tab, text="Tema Applicazione", font=("Arial", 18)).pack(pady=20)
        self.seg_theme = ctk.CTkSegmentedButton(tab, values=["Blue", "Green", "Dark-Blue"], command=self.change_theme)
        self.seg_theme.set("Blue")
        self.seg_theme.pack(pady=10)
        ctk.CTkLabel(tab, text="(Richiede riavvio per applicarsi completamente a tutti gli elementi)", text_color="gray").pack()

    def change_theme(self, value):
        ctk.set_default_color_theme(value.lower())
        msgbox.showinfo("Tema", f"Hai selezionato {value}. Riavviera l'app per vedere tutte le modifiche.")

    def ask_save(self):
        topic = self.entry.get().strip()
        fmt = self.opt_format.get()
        ext = ".pdf" if "PDF" in fmt else ".docx"
        
        path = filedialog.asksaveasfilename(defaultextension=ext, title="Salva Report", initialfile=f"{topic}{ext}")
        if path:
            self.toggle_ui(False)
            self.progress.pack()
            self.progress.start()
            threading.Thread(target=self.worker, args=(topic, self.opt_lang.get(), self.seg_depth.get(), path, fmt)).start()

    def worker(self, topic, lang, depth, path, fmt):
        try:
            generate_report(topic, lang, depth, path, fmt)
            self.saved_path = path
            self.on_success()
        except Exception as e:
            print(e)
            self.on_fail(str(e))

    def on_success(self):
        self.progress.stop()
        self.progress.pack_forget()
        self.status.configure(text="Completato!", text_color="green")
        self.btn_open.pack(pady=10)
        self.btn_open.configure(state="normal")
        self.toggle_ui(True)
        self.refresh_history()

    def on_fail(self, err):
        self.progress.stop()
        self.progress.pack_forget()
        self.status.configure(text=f"Errore: {err}", text_color="red")
        self.toggle_ui(True)
        
    def toggle_ui(self, enable):
        state = "normal" if enable else "disabled"
        self.btn_go.configure(state=state)
        
    def open_current(self):
        if self.saved_path: os.startfile(self.saved_path)

if __name__ == "__main__":
    app = UltimateApp()
    app.mainloop()
