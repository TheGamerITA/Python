import os
import re
from collections import Counter

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

from config import TODAY
from services.history_service import save_to_history
from services.search_service import collect_research


def clean_text(text):
    if not text:
        return ""
    replacements = {
        "€": "EUR",
        "”": '"',
        "“": '"',
        "’": "'",
        "‘": "'",
        "–": "-",
        "…": "...",
        "à": "a",
        "è": "e",
        "é": "e",
        "ì": "i",
        "ò": "o",
        "ù": "u",
        "Á": "A",
        "É": "E",
        "Í": "I",
        "Ó": "O",
        "Ú": "U",
        "ä": "a",
        "ö": "o",
        "ü": "u",
        "ß": "ss",
        "ñ": "n",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "replace").decode("latin-1")


def create_chart(text_data, topic):
    try:
        words = re.findall(r"\b\w{5,}\b", text_data.lower())
        ignore = {
            "anche", "della", "delle", "nella", "hanno", "stato", "sono", "come", "questo", "questa", "degli",
            "parte", "prima", "dopo", "tutto", "tutti", "fatto", "essere", "avere", "which", "their", "about",
            "would", "these", "other", "sur", "pour", "dans", "avec", "plus", "not", "that", "with", "from",
            "this", "have",
        }
        filtered = [w for w in words if w not in ignore]
        counts = Counter(filtered).most_common(7)
        if not counts:
            return None

        labels, values = zip(*counts)
        plt.figure(figsize=(6, 4))
        plt.bar(labels, values, color="#213363")
        plt.title(f"Keywords: {topic}")
        plt.xticks(rotation=45)
        plt.tight_layout()

        filename = "temp_chart.png"
        plt.savefig(filename)
        plt.close()
        return filename
    except Exception:
        return None


def generate_docx(topic, wiki_summary, web_results, img_file, chart_file, save_path):
    doc = Document()
    doc.add_heading(f"Report: {topic}", 0)
    doc.add_paragraph(f"Generated on: {TODAY()}")

    if img_file:
        try:
            doc.add_picture(img_file, width=Inches(4))
        except Exception:
            pass

    doc.add_heading("General Overview", level=1)
    doc.add_paragraph(wiki_summary)

    if chart_file:
        doc.add_heading("Data Analysis", level=1)
        try:
            doc.add_picture(chart_file, width=Inches(5))
        except Exception:
            pass

    if web_results:
        doc.add_heading("Web Resources", level=1)
        for res in web_results:
            p = doc.add_paragraph()
            runner = p.add_run(res["title"])
            runner.bold = True
            doc.add_paragraph(res["body"])
            doc.add_paragraph(res["href"], style="Intense Quote")

    doc.save(save_path)


class PDFReport(FPDF):
    def __init__(self, topic, cover_image=None):
        super().__init__()
        self.topic = clean_text(topic)
        self.cover_image = cover_image
        self.set_auto_page_break(auto=True, margin=15)

    def header(self):
        if self.page_no() > 1:
            self.set_font("Arial", "I", 8)
            self.set_text_color(128)
            self.cell(0, 10, f"Report: {self.topic}", 0, 0, "R")
            self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.set_text_color(128)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")

    def create_cover_page(self):
        self.add_page()
        self.set_y(40)
        self.set_font("Arial", "B", 24)
        self.set_text_color(33, 51, 99)
        self.cell(0, 20, "REPORT PREMIUM", ln=True, align="C")

        if self.cover_image:
            try:
                self.image(self.cover_image, x=65, y=70, w=80)
                self.set_y(160)
            except Exception:
                self.set_y(100)
        else:
            self.set_y(100)

        self.set_font("Arial", "B", 32)
        self.set_text_color(0)
        self.cell(0, 20, self.topic, ln=True, align="C")
        self.add_page()

    def add_section_title(self, title):
        self.set_font("Arial", "B", 14)
        self.set_text_color(33, 51, 99)
        self.cell(0, 10, clean_text(title), 0, 1, "L")
        self.set_text_color(0)
        self.ln(2)

    def add_paragraph(self, text):
        self.set_font("Arial", "", 11)
        self.multi_cell(0, 6, clean_text(text))
        self.ln(5)

    def add_web_card(self, title, body, link):
        self.set_fill_color(240, 240, 240)
        self.set_font("Arial", "B", 10)
        self.cell(0, 8, clean_text(title), 0, 1, "L", True)
        self.set_font("Arial", "", 10)
        self.multi_cell(0, 5, clean_text(body))
        self.set_font("Arial", "I", 8)
        self.set_text_color(0, 0, 255)
        self.cell(0, 5, clean_text(link), 0, 1)
        self.set_text_color(0)
        self.ln(3)


def generate_report(topic, lang, depth, save_path, export_format):
    wiki_summary, web_results, full_text, img_file = collect_research(topic, lang, depth)
    chart_file = create_chart(full_text, topic) if depth != "Fast" else None

    if "PDF" in export_format:
        pdf = PDFReport(topic, img_file)
        pdf.create_cover_page()
        pdf.add_section_title(f"Overview ({lang})")
        pdf.add_paragraph(wiki_summary)
        pdf.ln()
        if chart_file:
            pdf.add_section_title("Semantic Analysis")
            pdf.image(chart_file, x=50, w=110)
            pdf.ln(10)
        if web_results:
            pdf.add_section_title("Web Resources")
            for res in web_results:
                pdf.add_web_card(res["title"], res["body"], res["href"])
        pdf.output(save_path)
    elif "Word" in export_format:
        generate_docx(topic, wiki_summary, web_results, img_file, chart_file, save_path)

    if img_file and os.path.exists(img_file):
        os.remove(img_file)
    if chart_file and os.path.exists(chart_file):
        os.remove(chart_file)

    save_to_history(topic, save_path)
    return save_path
